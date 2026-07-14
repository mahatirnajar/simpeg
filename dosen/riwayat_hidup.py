"""
views_riwayat_hidup.py
======================
Generator Dokumen Riwayat Hidup (.docx) untuk Dosen dan Tendik UNTAD.

Struktur dokumen:
  - Kop surat (logo + nama instansi)
  - Garis ganda bawah kop
  - Judul DAFTAR RIWAYAT HIDUP
  - Jati Diri
  - Riwayat Pendidikan
  - Riwayat Kepangkatan
  - Riwayat Jabatan Fungsional
  - Riwayat Jabatan Struktural (jika ada)
  - Tugas Tambahan (jika ada)
  - Riwayat Keluarga
  - Tanda tangan

Setup:
  1. Letakkan logo di: BASE_DIR/static/img/logo_untad.png
  2. pip install python-docx

URL tambahkan di urls.py:
  # Tendik
  path('<int:pk>/riwayat-hidup/', views.riwayat_hidup_tendik, name='tendik_riwayat_hidup'),
  # Dosen (di dosen/urls.py)
  path('<int:pk>/riwayat-hidup/', views.riwayat_hidup_dosen, name='dosen_riwayat_hidup'),
"""

import io
import os
from datetime import date

from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.shortcuts import get_object_or_404

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import django.conf
from matplotlib import table
LOGO_PATH = os.path.join(
    django.conf.settings.BASE_DIR, 'static', 'logo', 'untad.png'
)

FONT_NAME = 'Times New Roman'

BULAN_ID = [
    '', 'Januari', 'Februari', 'Maret', 'April', 'Mei', 'Juni',
    'Juli', 'Agustus', 'September', 'Oktober', 'November', 'Desember',
]
HARI_ID = ['Senin', 'Selasa', 'Rabu', 'Kamis', 'Jumat', 'Sabtu', 'Minggu']

JENJANG_LABEL = {
    'SD': 'Sekolah Dasar (SD)',
    'SLTP': 'Sekolah Menengah Pertama (SMP)',
    'SLTA': 'Sekolah Menengah Atas (SMA)',
    'D3': 'Diploma III (D3)',
    'D4': 'Diploma IV (D4)',
    'S1': 'Sarjana (S1)',
    'S2': 'Magister (S2)',
    'S3': 'Doktor (S3)',
    'PROFESI': 'Profesi',
}


# ─────────────────────────────────────────────────────────────────────────────
# HELPER UMUM
# ─────────────────────────────────────────────────────────────────────────────

def _fmt_date(d):
    if not d:
        return '........'
    return f"{d.day} {BULAN_ID[d.month]} {d.year}"


def _fmt_date_hari(d):
    if not d:
        return '........'
    return f"{HARI_ID[d.weekday()]}, {d.day} {BULAN_ID[d.month]} {d.year}"


def _set_font(run, size=12, bold=False, italic=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def _fix_settings(doc):
    """Fix zoom percent yang wajib ada di settings.xml."""
    settings = doc.settings.element
    zoom = settings.find(qn('w:zoom'))
    if zoom is not None and not zoom.get(qn('w:percent')):
        zoom.set(qn('w:percent'), '100')


def _remove_tbl_borders(table):
    """Hapus semua border tabel."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    # Hapus tblBorders lama
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    # Buat baru
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    # Insert sebelum shd/tblLayout/tblCellMar/tblLook
    children = list(tblPr)
    insert_pos = len(children)
    for i, c in enumerate(children):
        if c.tag.split('}')[1] in ('shd', 'tblLayout', 'tblCellMar', 'tblLook'):
            insert_pos = i
            break
    tblPr.insert(insert_pos, tblBorders)


def _para(doc, text='', bold=False, italic=False, size=12,
          align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          space_before=0, space_after=6, left_indent=0):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if left_indent:
        pf.left_indent = Cm(left_indent)
    if text:
        run = p.add_run(text)
        _set_font(run, size=size, bold=bold, italic=italic)
    return p


def _add_run(para, text, bold=False, italic=False, size=12):
    run = para.add_run(text)
    _set_font(run, size=size, bold=bold, italic=italic)
    return run


def _section_title(doc, text):
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.LEFT,
              space_before=8, space_after=3)
    run = p.add_run(text)
    _set_font(run, size=12, bold=True)
    run.underline = True
    return p


def _cell_para(cell, text='', bold=False, size=12,
               align=WD_ALIGN_PARAGRAPH.CENTER):
    p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if text:
        run = p.add_run(text)
        _set_font(run, size=size, bold=bold)
    return p

def _set_col_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(int(width_cm * 567)))  # cm → twips (1 cm = 567 twips)
    tcW.set(qn('w:type'), 'dxa')


# ─────────────────────────────────────────────────────────────────────────────
# KOP SURAT
# ─────────────────────────────────────────────────────────────────────────────

def _buat_kop(doc):
    # Tabel 2 kolom: logo | teks
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    _remove_tbl_borders(table)

    _set_col_width(table.rows[0].cells[0], 3.5)   
    _set_col_width(table.rows[0].cells[1], 12.5)  

    # Logo
    cell_logo = table.rows[0].cells[0]
    cell_logo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.space_after = Pt(0)
    if os.path.exists(LOGO_PATH):
        p_logo.add_run().add_picture(LOGO_PATH, width=Cm(3.5))
    else:
        p_logo.add_run('[LOGO]')

    # Teks kop
    cell_txt = table.rows[0].cells[1]
    cell_txt.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Hapus paragraf kosong default
    for para in list(cell_txt.paragraphs):
        para._element.getparent().remove(para._element)

    for text, size, bold in [
        ('KEMENTERIAN PENDIDIKAN TINGGI,', 16, False),
        ('SAINS, DAN TEKNOLOGI', 16, False),
        ('UNIVERSITAS TADULAKO', 20, True),
        ('Jalan Soekarno Hatta Kilometer 9 Tondo, Mantikulore, Palu 94119', 12, False),
        ('Surel: untad@untad.ac.id   Laman: https://untad.ac.id', 12, False),
    ]:
        _cell_para(cell_txt, text, bold=bold, size=size)

    # Garis ganda bawah kop
    # Gunakan teknik: paragraf dengan border bottom via pPr
    # pBdr harus di posisi awal pPr (sebelum spacing/ind/jc)
    p_garis = doc.add_paragraph()
    p_garis.paragraph_format.space_before = Pt(2)
    p_garis.paragraph_format.space_after = Pt(4)
    pPr = p_garis._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'double')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '000000')
    pBdr.append(bot)
    pPr.insert(0, pBdr)  # insert di posisi 0 → sebelum semua elemen lain


# ─────────────────────────────────────────────────────────────────────────────
# BLOK TANDA TANGAN
# ─────────────────────────────────────────────────────────────────────────────

def _buat_ttd(doc, today):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.allow_autofit = False
    _remove_tbl_borders(table)
    _set_col_width(table.rows[0].cells[0], 8.0)
    _set_col_width(table.rows[0].cells[1], 8.5)

    cell_ttd = table.rows[0].cells[1]
    for para in list(cell_ttd.paragraphs):
        para._element.getparent().remove(para._element)

    for text, bold in [
        (f'DI TETAPKAN DI  : PALU', False),
        (f'PADA TANGGAL    : {_fmt_date(today)}', False),
        ('', False),
        ('an. Rektor', False),
        ('Wakil Rektor Bidang Keuangan dan Umum', False),
        ('', False),
        ('', False),
        ('', False),
        ('', False),
        ('Prof. Dr. M. Rusydi H, M.Si.', True),
        ('NIP. 196311131992031001', False),
    ]:
        _cell_para(cell_ttd, text, bold=bold, size=12)


# ─────────────────────────────────────────────────────────────────────────────
# GENERATOR TENDIK
# ─────────────────────────────────────────────────────────────────────────────

def _buat_riwayat_tendik(tendik):
    doc = Document()

    # Margin
    sec = doc.sections[0]
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.0)

    # Fix settings.xml
    _fix_settings(doc)

    # ── Kop ──────────────────────────────────────────────────────────────────
    _buat_kop(doc)

    # ── Cek kondisi meninggal ─────────────────────────────────────────────────
    riwayat_berhenti = getattr(tendik, 'riwayat_berhenti', None)
    meninggal = (
        riwayat_berhenti is not None and
        hasattr(riwayat_berhenti, 'alasan') and
        riwayat_berhenti.alasan == 'MENINGGAL'
    )

    jk = tendik.jenis_kelamin
    if meninggal:
        alm = 'ALMARHUM' if jk == 'L' else 'ALMARHUMAH'
        alm_kecil = 'Almarhum' if jk == 'L' else 'Almarhumah'
        judul = f'DAFTAR RIWAYAT HIDUP {alm}'
        sapaan = alm_kecil
        # Basmalah (Islam)
        if tendik.agama == 'Islam':
            p_arab = _para(doc, 'بِسْمِ اللَّهِ الرَّحْمَنِ الرَّحِيمِ',
                           align=WD_ALIGN_PARAGRAPH.CENTER,
                           size=16, space_before=4, space_after=0)
            _para(doc, 'INNALILLAHI WAINNA ILAIHI ROJIUN',
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=12, space_before=0, space_after=4)
    else:
        judul = 'DAFTAR RIWAYAT HIDUP'
        sapaan = 'Yang bersangkutan'

    # ── Judul ─────────────────────────────────────────────────────────────────
    _para(doc, space_before=4, space_after=0)
    _para(doc, judul, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
          size=14, space_before=0, space_after=2)
    _para(doc, tendik.nama_lengkap, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, size=14, space_before=0, space_after=8)

    # ── JATI DIRI ─────────────────────────────────────────────────────────────
    _section_title(doc, 'JATI DIRI')

    kp   = tendik.kepangkatan_terakhir
    jf   = tendik.jabatan_fungsional_terakhir
    js   = tendik.jabatan_struktural_aktif

    pangkat_gol = (
        f"{kp.pangkat} Gol. {kp.golongan}" if kp else '........'
    )
    jabatan_str = ''
    if jf:
        jabatan_str = f" Jabatan {jf.nama_dan_jenjang}."
    elif js:
        jabatan_str = f" Jabatan {js.jabatan}."

    unit_str = tendik.unit_kerja.nama if tendik.unit_kerja else '........'

    p_jati = _para(doc, space_before=0, space_after=6,
                   align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    _add_run(p_jati,
        f"{sapaan} {tendik.nama_lengkap}, "
        f"NIP. {tendik.nip or '........'}, "
        f"Lahir di {tendik.tempat_lahir or '........'} "
        f"tanggal {_fmt_date(tendik.tanggal_lahir)}. "
        f"Beragama {tendik.agama or '........'}. "
        f"Pangkat/Golongan {pangkat_gol}.{jabatan_str} "
        f"Unit Kerja {unit_str}."
    )

    # ── RIWAYAT PENDIDIKAN ────────────────────────────────────────────────────
    _section_title(doc, 'RIWAYAT PENDIDIKAN')
    _para(doc, 'Pendidikan formal :', size=12, space_before=0, space_after=2)

    pendidikan = list(tendik.riwayat_pendidikan.order_by('jenjang').all())
    if pendidikan:
        for pend in pendidikan:
            label = JENJANG_LABEL.get(pend.jenjang, pend.jenjang)
            prodi = (f" Jurusan/Prodi {pend.bidang_studi}"
                     if pend.bidang_studi else '')
            lulus = (f" pada Tahun {pend.tahun_lulus}"
                     if pend.tahun_lulus else '')
            _para(doc,
                  f"Lulus {label}{prodi} pada {pend.institusi or '........'}{lulus};",
                  size=12, space_before=0, space_after=2, left_indent=0.5)
    else:
        _para(doc, '(Belum ada data riwayat pendidikan)',
              size=12, space_before=0, space_after=2, left_indent=0.5)

    # ── RIWAYAT KEPANGKATAN ───────────────────────────────────────────────────
    _section_title(doc, 'RIWAYAT KEPANGKATAN')

    kepangkatan = list(tendik.riwayat_kepangkatan.order_by('tmt').all())
    if kepangkatan:
        for i, kp in enumerate(kepangkatan):
            if i == 0:
                teks = (
                    f"Diangkat sebagai Calon Pegawai Negeri Sipil"
                    f"{f' dalam pangkat {kp.pangkat}' if kp.pangkat else ''}"
                    f" Gol. {kp.golongan},"
                    f" terhitung mulai tanggal {_fmt_date(kp.tmt)}."
                )
            else:
                teks = (
                    f"Diangkat sebagai"
                    f"{f' {kp.pangkat}' if kp.pangkat else ' ........'}"
                    f" Gol. {kp.golongan},"
                    f" terhitung mulai tanggal {_fmt_date(kp.tmt)}."
                )
            _para(doc, teks, size=12, space_before=0, space_after=2,
                  left_indent=0.5)
    else:
        _para(doc, '(Belum ada data riwayat kepangkatan)',
              size=12, space_before=0, space_after=2, left_indent=0.5)

    # ── RIWAYAT JABATAN FUNGSIONAL ────────────────────────────────────────────
    _section_title(doc, 'RIWAYAT JABATAN FUNGSIONAL')

    jabatan_fungsional = list(
        tendik.riwayat_jabatan_fungsional.order_by('tmt').all()
    )
    if jabatan_fungsional:
        for jf in jabatan_fungsional:
            nama = jf.nama_dan_jenjang or jf.nama_jabatan or '........'
            _para(doc,
                  f"Diangkat sebagai {nama},"
                  f" terhitung mulai tanggal {_fmt_date(jf.tmt)}.",
                  size=12, space_before=0, space_after=2, left_indent=0.5)
    else:
        _para(doc, '(Belum ada data riwayat jabatan fungsional)',
              size=12, space_before=0, space_after=2, left_indent=0.5)

    # ── RIWAYAT JABATAN STRUKTURAL (jika ada) ─────────────────────────────────
    jabatan_struktural = list(
        tendik.jabatan_struktural.order_by('tmt_jabatan').all()
    )
    if jabatan_struktural:
        _section_title(doc, 'RIWAYAT JABATAN STRUKTURAL')
        for js in jabatan_struktural:
            eselon = f" (Eselon {js.eselon})" if js.eselon else ''
            _para(doc,
                  f"Diangkat sebagai {js.jabatan}{eselon},"
                  f" terhitung mulai tanggal {_fmt_date(js.tmt_jabatan)}.",
                  size=12, space_before=0, space_after=2, left_indent=0.5)

    # ── TUGAS TAMBAHAN (jika ada) ─────────────────────────────────────────────
    tugas_tambahan = list(tendik.tugas_tambahan.order_by('tmt_jabatan').all())
    if tugas_tambahan:
        _section_title(doc, 'TUGAS TAMBAHAN')
        for tt in tugas_tambahan:
            aktif = '' if tt.is_aktif else ' (s.d. ......)'
            _para(doc,
                  f"Diangkat sebagai {tt.jabatan}{aktif},"
                  f" terhitung mulai tanggal {_fmt_date(tt.tmt_jabatan)}.",
                  size=12, space_before=0, space_after=2, left_indent=0.5)

    # ── RIWAYAT KELUARGA ──────────────────────────────────────────────────────
    _section_title(doc, 'RIWAYAT KELUARGA')

    if jk == 'L':
        _para(doc, f'{sapaan} menikahi seorang istri yang bernama :',
              size=12, space_before=0, space_after=2, left_indent=0.5)
    else:
        _para(doc, f'{sapaan} menikah dengan seorang suami yang bernama :',
              size=12, space_before=0, space_after=2, left_indent=0.5)
    _para(doc, '1.  .......................................',
          size=12, space_before=0, space_after=2, left_indent=1.0)
    _para(doc, 'Dari pernikahan tersebut dikaruniai ....... orang anak :',
          size=12, space_before=0, space_after=2, left_indent=0.5)
    _para(doc, '1.  .......................................',
          size=12, space_before=0, space_after=2, left_indent=1.0)

    # Kalimat meninggal (jika berlaku)
    if meninggal and riwayat_berhenti:
        tgl_str = _fmt_date_hari(riwayat_berhenti.tanggal)
        ket = f" di {riwayat_berhenti.keterangan}" if riwayat_berhenti.keterangan else ''
        _para(doc,
              f"{sapaan} meninggal dunia pada {tgl_str},"
              f" pukul ..... WITA{ket}.",
              size=12, space_before=4, space_after=6,
              align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # ── Penutup & TTD ────────────────────────────────────────────────────────
    _para(doc, 'Demikian Riwayat Hidup ini dibuat.',
          size=12, space_before=8, space_after=12)
    _buat_ttd(doc, date.today())

    return doc


# ─────────────────────────────────────────────────────────────────────────────
# GENERATOR DOSEN
# ─────────────────────────────────────────────────────────────────────────────

def _buat_riwayat_dosen(dosen):
    doc = Document()

    sec = doc.sections[0]
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.0)

    _fix_settings(doc)
    _buat_kop(doc)


    # ── Judul ─────────────────────────────────────────────────────────────────
    # ── Cek kondisi meninggal ─────────────────────────────────────────────────
    status_terakhir = dosen.status_terakhir
    meninggal = status_terakhir and status_terakhir.status == 'MENINGGAL'
    jk = dosen.jenis_kelamin
    if meninggal:
        alm = 'ALMARHUM' if jk == 'L' else 'ALMARHUMAH'
        alm_kecil = 'Almarhum' if jk == 'L' else 'Almarhumah'
        judul = f'DAFTAR RIWAYAT HIDUP {alm}'
        sapaan = alm_kecil
        # Basmalah (Islam)
        if dosen.agama == 'Islam':
            p_arab = _para(doc, 'بِسْمِ اللَّهِ الرَّحْمَنِ الرَّحِيمِ',
                           align=WD_ALIGN_PARAGRAPH.CENTER,
                           size=16, space_before=4, space_after=0)
            _para(doc, 'INNALILLAHI WAINNA ILAIHI ROJIUN',
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                  size=12, space_before=0, space_after=4)
    else:
        judul = 'DAFTAR RIWAYAT HIDUP'
        sapaan = 'Yang bersangkutan'
        
    # ── Judul ─────────────────────────────────────────────────────────────────
    _para(doc, space_before=4, space_after=0)
    _para(doc, judul, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
          size=14, space_before=0, space_after=2)
    _para(doc, dosen.nama_lengkap, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, size=14, space_before=0, space_after=8)

    # ── JATI DIRI ─────────────────────────────────────────────────────────────
    _section_title(doc, 'JATI DIRI')

    kp  = dosen.kepangkatan_terakhir
    jf  = dosen.jabatan_fungsional_terakhir
    tt  = dosen.tugas_tambahan_aktif

    pangkat_gol = (
        f"{kp.pangkat} Gol. {kp.golongan}" if kp else '........'
    )
    jabatan_str = f" Jabatan {jf.jenjang}." if jf else ''
    tugas_str   = f" Tugas Tambahan sebagai {tt.jabatan}." if tt else ''
    fak_str     = dosen.fakultas.nama if dosen.fakultas else '........'
    prodi_str   = dosen.program_studi_nama or '........'

    p_jati = _para(doc, space_before=0, space_after=6,
                   align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    _add_run(p_jati,
        f"{sapaan} {dosen.nama_lengkap}, "
        f"NIP. {dosen.nip or '........'}, "
        f"Lahir di {dosen.tempat_lahir or '........'} "
        f"tanggal {_fmt_date(dosen.tanggal_lahir)}. "
        f"Beragama {dosen.agama or '........'}. "
        f"Pangkat/Golongan {pangkat_gol}.{jabatan_str}{tugas_str} "
        f"Fakultas {fak_str}, Program Studi {prodi_str}."
    )

    # ── RIWAYAT PENDIDIKAN ────────────────────────────────────────────────────
    _section_title(doc, 'RIWAYAT PENDIDIKAN')
    _para(doc, 'Pendidikan formal :', size=12, space_before=0, space_after=2)

    pendidikan = list(dosen.riwayat_pendidikan.order_by('jenjang').all())
    if pendidikan:
        for pend in pendidikan:
            label = JENJANG_LABEL.get(pend.jenjang, pend.jenjang)
            prodi = (f" Jurusan/Prodi {pend.bidang_studi}"
                     if pend.bidang_studi else '')
            lulus = (f" pada Tahun {pend.tahun_lulus}"
                     if pend.tahun_lulus else '')
            _para(doc,
                  f"Lulus {label}{prodi} pada {pend.institusi or '........'}{lulus};",
                  size=12, space_before=0, space_after=2, left_indent=0.5)
    else:
        _para(doc, '(Belum ada data riwayat pendidikan)',
              size=12, space_before=0, space_after=2, left_indent=0.5)

    # ── RIWAYAT KEPANGKATAN ───────────────────────────────────────────────────
    _section_title(doc, 'RIWAYAT KEPANGKATAN')

    kepangkatan = list(dosen.riwayat_kepangkatan.order_by('tmt').all())
    if kepangkatan:
        for i, kp in enumerate(kepangkatan):
            if i == 0:
                teks = (
                    f"{sapaan} Diangkat sebagai Calon Pegawai Negeri Sipil"
                    f" dalam pangkat {kp.pangkat or '........'} Gol. {kp.golongan},"
                    f" terhitung mulai tanggal {_fmt_date(kp.tmt)}."
                )
            else:
                teks = (
                    f"{sapaan} Diangkat sebagai {kp.pangkat or '........'}"
                    f" Gol. {kp.golongan},"
                    f" terhitung mulai tanggal {_fmt_date(kp.tmt)}."
                )
            _para(doc, teks, size=12, space_before=0, space_after=2,
                  left_indent=0.5)
    else:
        _para(doc, '(Belum ada data riwayat kepangkatan)',
              size=12, space_before=0, space_after=2, left_indent=0.5)

    # ── RIWAYAT JABATAN FUNGSIONAL ────────────────────────────────────────────
    _section_title(doc, 'RIWAYAT JABATAN FUNGSIONAL')

    jabatan = list(dosen.riwayat_jabatan_fungsional.order_by('tmt').all())
    if jabatan:
        for jf in jabatan:
            _para(doc,
                  f"{sapaan} Diangkat sebagai {jf.jenjang or '........'},"
                  f" terhitung mulai tanggal {_fmt_date(jf.tmt)}.",
                  size=12, space_before=0, space_after=2, left_indent=0.5)
    else:
        _para(doc, '(Belum ada data riwayat jabatan fungsional)',
              size=12, space_before=0, space_after=2, left_indent=0.5)

    # ── TUGAS TAMBAHAN (jika ada) ─────────────────────────────────────────────
    tugas = list(dosen.tugas_tambahan.order_by('tmt_jabatan').all())
    if tugas:
        _section_title(doc, 'TUGAS TAMBAHAN')
        for tt in tugas:
            aktif = '' if tt.is_aktif else ' (s.d. ......)'
            _para(doc,
                  f"{sapaan} Diangkat sebagai {tt.jabatan}{aktif},"
                  f" terhitung mulai tanggal {_fmt_date(tt.tmt_jabatan)}.",
                  size=12, space_before=0, space_after=2, left_indent=0.5)

    # ── RIWAYAT KELUARGA ──────────────────────────────────────────────────────
    _section_title(doc, 'RIWAYAT KELUARGA')

    if dosen.jenis_kelamin == 'L':
        _para(doc, f"{sapaan} menikahi seorang istri yang bernama :",
              size=12, space_before=0, space_after=2, left_indent=0.5)
    else:
        _para(doc, f"{sapaan} menikah dengan seorang suami yang bernama :",
              size=12, space_before=0, space_after=2, left_indent=0.5)
    _para(doc, '1.  .......................................',
          size=12, space_before=0, space_after=2, left_indent=1.0)
    _para(doc, 'Dari pernikahan tersebut dikaruniai ....... orang anak :',
          size=12, space_before=0, space_after=2, left_indent=0.5)
    _para(doc, '1.  .......................................',
          size=12, space_before=0, space_after=2, left_indent=1.0)

    _para(doc, 'Demikian Riwayat Hidup ini dibuat.',
          size=12, space_before=8, space_after=12)
    _buat_ttd(doc, date.today())

    return doc


# ─────────────────────────────────────────────────────────────────────────────
# DJANGO VIEWS
# ─────────────────────────────────────────────────────────────────────────────

@login_required
def riwayat_hidup_tendik(request, pk):
    from tendik.models import Tendik
    tendik = get_object_or_404(
        Tendik.objects.select_related('unit_kerja').prefetch_related(
            'riwayat_kepangkatan',
            'riwayat_jabatan_fungsional',
            'jabatan_struktural',
            'riwayat_pendidikan',
            'tugas_tambahan',
            'riwayat_berhenti',
        ),
        pk=pk,
    )
    doc = _buat_riwayat_tendik(tendik)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    nama_file = f"Riwayat_Hidup_{tendik.nama_terang.replace(' ', '_')}.docx"
    resp = HttpResponse(
        buf.getvalue(),
        content_type=(
            'application/vnd.openxmlformats-officedocument'
            '.wordprocessingml.document'
        ),
    )
    resp['Content-Disposition'] = f'attachment; filename="{nama_file}"'
    return resp


@login_required
def riwayat_hidup_dosen(request, pk):
    from dosen.models import Dosen
    dosen = get_object_or_404(
        Dosen.objects.select_related('fakultas').prefetch_related(
            'riwayat_kepangkatan',
            'riwayat_jabatan_fungsional',
            'riwayat_pendidikan',
            'tugas_tambahan',
        ),
        pk=pk,
    )
    doc = _buat_riwayat_dosen(dosen)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    nama_file = f"Riwayat_Hidup_{dosen.nama_terang.replace(' ', '_')}.docx"
    resp = HttpResponse(
        buf.getvalue(),
        content_type=(
            'application/vnd.openxmlformats-officedocument'
            '.wordprocessingml.document'
        ),
    )
    resp['Content-Disposition'] = f'attachment; filename="{nama_file}"'
    return resp